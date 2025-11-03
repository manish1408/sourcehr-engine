import os
import random
from dotenv import load_dotenv
import time
from datetime import datetime
import tempfile
from urllib.parse import urlparse
import os
import requests
import docx  # python-docx
from langchain_openai import AzureChatOpenAI, OpenAIEmbeddings, AzureOpenAIEmbeddings
from pinecone import Pinecone, ServerlessSpec, Index
from pinecone.exceptions import NotFoundException
from app.helpers.Utilities import Utils
from langchain_pinecone import PineconeVectorStore
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownTextSplitter
from uuid import uuid4
from langchain_core.documents import Document
from dotenv import load_dotenv
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.schema import HumanMessage, SystemMessage

from app.schemas.Dashboard import  LawChangeListByLocation, NewsList,LegalCalendar, CourtDecisionList
from app.schemas.MetaDataSchema import MetaDataSchemaList
from app.helpers.MetaDataHelper import MetaDataHelper
from app.models.Locations import LocationsModel
from app.models.Industries import IndustriesModel
from app.models.Topics import TopicsModel
load_dotenv()

class VectorDB:
    def __init__(self,namespace):
        pc = Pinecone(api_key=os.getenv('PINECONE_API_KEY'))
        index_name = os.getenv("PINECONE_INDEX")
        # embeddings = OpenAIEmbeddings(model="text-embedding-3-large", api_key=os.getenv('OPENAI_API_KEY'))
        self.embeddings = AzureOpenAIEmbeddings(
            model="text-embedding-3-large",
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
            api_key=os.getenv('AZURE_OPENAI_KEY'),
            openai_api_version='2024-12-01-preview'
        )
        
        self.index = pc.Index(index_name)
        self.namespace = namespace
        self.vector_store = PineconeVectorStore(index=self.index, embedding=self.embeddings,namespace=namespace)
        self.meta_data_helper = MetaDataHelper()
        self.locations_model=LocationsModel()
        self.industries_model=IndustriesModel()
        self.topics_model=TopicsModel()
        
        
    def generate_embedding(self, text: str):
        # response = self.azure_embedding_client.embeddings.create(
        #     model="text-embedding-3-large",
        #     input=text
        # )
        
        response = self.embeddings.embed_query(text)
        return response 

    def retrieve_by_metadata(self, query: str, metadata_filter: dict, k: int = 5):
            """
            Retrieve documents based on query + metadata filtering.
            """
            try:
                query_vector = self.generate_embedding(query)
                response = self.index.query(
                    vector=query_vector,
                    top_k=k,
                    filter=metadata_filter,
                    namespace=self.namespace,
                    include_metadata=True
                )

                # Convert results to LangChain Document objects
                docs = [
                    Document(
                        page_content=match['metadata'].get("text", ""),
                        metadata=match['metadata'],
                        id = match['id']
                    )
                    for match in response['matches']
                ]
                return docs
            except Exception as e:
                print(f"Error in retrieve_by_metadata: {e}")
                return []
                
    def enterDocumentToKnowledge(self, file_path, file_name='Untitled document',source_type=''):
        try:
            utc_now = datetime.utcnow()
            utc_string = utc_now.strftime("%Y-%m-%dT%H:%M:%S.%fZ")

            # Ensure we have a local file path. If a URL is provided, download to a temp file first.
            local_path = file_path
            is_temp_file = False
            try:
                if isinstance(file_path, str) and file_path.lower().startswith(("http://", "https://")):
                    parsed = urlparse(file_path)
                    _, ext = os.path.splitext(parsed.path)
                    fd, tmp_path = tempfile.mkstemp(suffix=ext or "")
                    os.close(fd)
                    resp = requests.get(file_path, timeout=60)
                    resp.raise_for_status()
                    with open(tmp_path, "wb") as f:
                        f.write(resp.content)
                    local_path = tmp_path
                    is_temp_file = True

                _, ext = os.path.splitext(local_path)
                ext = (ext or "").lower()

                documents = None
                if ext == ".pdf":
                    loader = PyPDFLoader(local_path)
                    documents = loader.load()
                elif ext == ".txt":
                    with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    documents = [Document(page_content=text, metadata={})]
                elif ext == ".docx":
                    try:
                        doc = docx.Document(local_path)
                        text = "\n".join([p.text for p in doc.paragraphs])
                        documents = [Document(page_content=text, metadata={})]
                    except Exception as e:
                        raise RuntimeError("DOCX support requires 'python-docx' to be installed.") from e
                else:
                    # Fallback: attempt to read as UTF-8 text
                    try:
                        with open(local_path, "r", encoding="utf-8") as f:
                            text = f.read()
                        documents = [Document(page_content=text, metadata={})]
                    except Exception:
                        raise ValueError(f"Unsupported file type for knowledge ingestion: '{ext}'")
            finally:
                if is_temp_file and os.path.exists(local_path):
                    try:
                        os.remove(local_path)
                    except Exception:
                        pass

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=20000,
                chunk_overlap=200,
            )
            chunks = text_splitter.split_documents(documents)

            results = []
            for chunk in chunks:
                try:
                    event_chunk_list = self.extract_meta_data_from_chunk(chunk.page_content)
                    for event_chunk in event_chunk_list.metaData:
                        event_chunk_uuid = str(uuid4())
                        event_chunk_document = Document(
                            id=event_chunk_uuid,
                            page_content=event_chunk.chunkText,
                            metadata={
                                "region_slug": event_chunk.region_slug,
                                "location_slug": event_chunk.location_slug,
                                "primary_industry_slug": event_chunk.primary_industry_slug,
                                "secondary_industry_slug": event_chunk.secondary_industry_slug,
                                "topic_slug": event_chunk.topic_slug,
                                "newsPublishTimestamp": datetime.fromisoformat(event_chunk.newsPublishTimestamp).timestamp(),
                                "discussedTimestamp": datetime.fromisoformat(event_chunk.discussedTimestamp).timestamp(),
                                "file_name": file_name,
                                "file_url": file_path,
                                 "sourceType": source_type,
                                "ingestion_date": utc_string,
                            }
                        )
                        try:
                            self.vector_store.add_documents(documents=[event_chunk_document], ids=[event_chunk_uuid], namespace=self.namespace)
                        except Exception as e:
                            print(f"Error adding document to vector store: {e}")
                        results.append({
                            "uuid": event_chunk_uuid,
                            "metadata": event_chunk_document.metadata,
                        })
                except Exception as e:
                    print(f"Metadata extraction failed for chunk: {e}")
            return results
        except Exception as e:
            print(f"Error in enterDocumentToKnowledge: {e}")
            return None
            
    def enterTextKnowledge(self, knowledge_text,file_name = 'Untitled document'):
        try:
            utc_now = datetime.utcnow()
            utc_string = utc_now.strftime("%Y-%m-%dT%H:%M:%S.%fZ") 

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=20000,
                chunk_overlap=200
            )
            chunks = text_splitter.split_text(knowledge_text)

            documents = []
            uuids = []

            for idx, chunk in enumerate(chunks):
                doc_id = str(uuid4())
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "chunk_index": idx,
                        "original_text_length": len(knowledge_text),
                        "file_name":file_name,
                        "ingestion_date": utc_string,
                    }
                )
                documents.append(doc)
                uuids.append(doc_id)

            self.vector_store.add_documents(documents=documents, ids=uuids, namespace=self.namespace)
            return uuids
        except Exception as e:
            return None
        
    
        
    def enterWebsiteToKnowledge(self, page_content, url, limit=20, source_type: str = ""):
        try:
            utc_now = datetime.utcnow()
            utc_string = utc_now.strftime("%Y-%m-%dT%H:%M:%S.%fZ") 
            documents = []
            doc = Document(
                page_content=page_content,
                metadata={
                    "pageUrl": url,
                    "sourceType": source_type,
                    "ingestion_date": utc_string,
                },
            )
            documents.append(doc)

            chunks = [doc]

            for chunk in chunks:
                results = []
                try:
                        event_chunk_list = self.extract_meta_data_from_chunk(chunk.page_content)
                        for event_chunk in event_chunk_list.metaData:
                            # self.meta_data_helper.check_if_new_data(event_chunk)
                            event_chunk_uuid = str(uuid4())

                            event_chunk_document = Document(
                            id=event_chunk_uuid,
                            page_content=event_chunk.chunkText,
                            metadata={
                                "region_slug": event_chunk.region_slug,
                                "location_slug": event_chunk.location_slug,
                                "primary_industry_slug": event_chunk.primary_industry_slug,
                                "secondary_industry_slug": event_chunk.secondary_industry_slug,                   
                                "topic_slug": event_chunk.topic_slug,
                                "newsPublishTimestamp": datetime.fromisoformat(event_chunk.newsPublishTimestamp).timestamp(),
                                "discussedTimestamp": datetime.fromisoformat(event_chunk.discussedTimestamp).timestamp(),
                                "pageUrl": url,
                                "sourceType": source_type,
                                "ingestion_date": utc_string,
                            }
                        )
                            
                            try:
                                ingested=self.vector_store.add_documents(documents=[event_chunk_document], ids=[event_chunk_uuid], namespace=self.namespace)
                            except Exception as e:
                                print(f"Error adding document to vector store: {e}")
                            results.append({
                            "uuid": event_chunk_uuid,
                            "metadata": event_chunk_document.metadata,
                            })

                except Exception as e:
                    print(f"Metadata extraction failed for chunk: {e}")

            return results
        except Exception as e:
            print(f"Error in enterWebsiteToKnowledge: {e}")
            return None
        
    # def enterWebsiteToKnowledge(self, page_content,url,limit=20):
    #     try:
    #         utc_now = datetime.utcnow()
    #         utc_string = utc_now.strftime("%Y-%m-%dT%H:%M:%S.%fZ") 
    #         # pages = self.__crawler.crawl_a_webpage(url,limit)
    #         documents = []
    #         doc = Document(
    #                 page_content=page_content,
    #                 metadata={
    #                     "pageUrl": url,
    #                     "ingestion_date": utc_string,
    #                 },
    #             )
    #         documents.append(doc)

    #         text_splitter = MarkdownTextSplitter(
    #             chunk_size=20000,
    #             chunk_overlap=200,
    #         )
    #         chunks = text_splitter.split_documents(documents)
                
    #         uuids = [str(uuid4()) for _ in range(len(chunks))]

    #         self.vector_store.add_documents(documents=chunks, ids=uuids, namespace=self.namespace)
    #         return uuids
    
    #     except Exception as e:
    #         return None
        
    def deleteDocument(self, ids):
        try:
            # Support both list of dicts (with 'uuid') and list of strings
            if ids and isinstance(ids[0], dict) and 'uuid' in ids[0]:
                id_list = [item['uuid'] for item in ids]
            else:
                id_list = ids
            resp = self.vector_store.delete(id_list)
            return True
        except Exception as e:
            print(e)
            return False
    
    def deleteNamespace(self,namespace):
        try:
            resp = self.vector_store.delete(delete_all=True,namespace=namespace)
            return True
        except NotFoundException as e:
            print(f"Namespace '{namespace}' not found: {e}")
            return True
        except Exception as e:
            return False
        
    def get_vector_retriever(self):
        retriever = self.vector_store.as_retriever(
            search_type="similarity_score_threshold",
            search_kwargs={"k": 3, "score_threshold": 0.5}        )
        return retriever
    
    def extract_meta_data_from_chunk(self, scrapped_data: str):
        llm=AzureChatOpenAI(
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
            azure_deployment=os.getenv('gpt-4o-mini'),
            api_key=os.getenv('AZURE_OPENAI_KEY'),
            api_version="2024-12-01-preview",
            model_name="gpt-4o-mini",
            openai_api_type="azure",
            streaming=True,
            callbacks=[StreamingStdOutCallbackHandler()]
        )
        
        
        regions=list(self.locations_model.collection.find({}, {'region_name': 1, 'region_slug': 1, 'locations': 1}))
        region_lines = []
        for r in regions:
            locations = ", ".join([loc['name'] for loc in r['locations']])
            region_lines.append(f"- {r['region_name']} (Slug: {r['region_slug']}) → [{locations}]")
            
        industries=list(self.industries_model.collection.find({}))
        industry_lines = []
        for industry in industries:
            primary = industry["primary_industry"]
            primary_slug = industry["primary_industry_slug"]
            secondaries = ", ".join([
                f"{s['name']} (Slug: {s['slug']})" for s in industry.get("secondary_industry", [])
            ])
            line = f"- {primary} (Slug: {primary_slug}) → [{secondaries}]"
            industry_lines.append(line)
        topics=list(self.topics_model.collection.find({}))
        topic_lines = []
        for topic_group in topics:
            category = topic_group["category"]
            category_slug = topic_group["category_slug"]
            titles = ", ".join([
                f"{t['title']} (Slug: {t['slug']})" for t in topic_group.get("topics", [])
            ])
            line = f"- {category} (Slug: {category_slug}) → [{titles}]"
            topic_lines.append(line)
        


        
        now = datetime.utcnow().isoformat() + "Z"
        system_message = SystemMessage(
            content=f"""
        You are a legal metadata extraction engine.

        Current UTC timestamp: {now}

        You will receive a chunk of HR, employment law, business policy, or court decision text.
        - Extract ALL HR or employment-related events, including:
            - Legislative updates
            - Workplace policy changes
            - Court decisions, rulings, or judgments
            - Compliance mandates
            - Regulatory announcements
        
        ### Valid Regions:
        {chr(10).join(region_lines)}

        ### Valid Industries:
        {chr(10).join(industry_lines)}

        ### Valid Topics:
        {chr(10).join(topic_lines)}

        ### Guidelines:
        - Extract all HR-relevant events
        - Each event must be returned as a separate object with the following fields:
            - `chunkText`: A raw copy of the sentence(s) directly describing the event (no paraphrasing or shortening).
            - `region`: The geopolitical subregion where the event occurs, e.g., "North America", "Europe", "Asia-Pacific".
            - `region_slug`: A slugified version of the region.
            - `location`: The specific location  if applicable (e.g., "Bavaria"). If no state is mentioned return "".
            - `location_slug`: The specific location or province if applicable (e.g., "Bavaria"). If no state is mentioned return "".
            - `primary_industry`: The broad industry most impacted (e.g., "Finance", "Technology", "Healthcare").
            - `primary_industry__slug`: A slugified version of the primary industry (e.g., "finance", "technology", "healthcare").
            - `secondary_industry`: A more specific sector if applicable (e.g., "Fintech", "Pharmaceuticals").
            - `secondary_industry__slug`: A slugified version of the secondary industry (e.g., "fintech", "pharmaceuticals").
            - `topic`: The HR-relevant topic of the event (e.g., "Workplace Policy", "Discrimination", "Employment Law").
            - `topic_slug`: A slugified version of the topic (e.g., "workplace-policy", "discrimination", "employment-law").
            - `discussedTimestamp`: ISO 8601 date when the event occurred or will occur.
            - `newsPublishTimestamp`: ISO 8601 timestamp when the news article was published.
            
        """
        )


        messages = [
            system_message,
            HumanMessage(content=f"Chunk: {scrapped_data}")
        ]

        structured_llm = llm.with_structured_output(MetaDataSchemaList)
        response = structured_llm.invoke(messages)
        return response
        
    
    def extract_law_changes(self, vectors: list,location_slugs,industry_slugs,topic_slugs):
        vector_list = []

        for match in vectors.get("matches", []):
            text = match["metadata"].get("chunk_text") or match["metadata"].get("text", "")
            metadata = match["metadata"]
            vector_list.append({
                "text": text,
                "metadata": metadata
            })
        vectors_as_string = "\n\n---\n\n".join([
            f"Text: {v['text']}\n\nMetadata: {v['metadata']}" for v in vector_list
        ])
        llm=AzureChatOpenAI(
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
            azure_deployment=os.getenv('gpt-4o-mini'),
            api_key=os.getenv('AZURE_OPENAI_KEY'),
            api_version="2024-12-01-preview",
            model_name="gpt-4o-mini",
            openai_api_type="azure"        )
        
        location_names = [slug.replace('-', ' ').title() for slug in location_slugs]
        industry_names = [slug.replace('-', ' ').title() for slug in industry_slugs]
        topic_names = [slug.replace('-', ' ').title() for slug in topic_slugs]
        
        system_message = SystemMessage(
            content=f"""CRITICAL: You must STRICTLY follow these filtering rules:

        REQUIRED LOCATIONS ONLY: {', '.join(location_names) if location_names else 'None specified'}
        REQUIRED INDUSTRIES ONLY: {', '.join(industry_names) if industry_names else 'None specified'}  
        REQUIRED TOPICS ONLY: {', '.join(topic_names) if topic_names else 'None specified'}

        STRICT RULES:
        1. ONLY return law changes for the locations listed above - NO OTHER jurisdictions
        2. ONLY include law changes related to the specified industries
        3. ONLY include law changes about the specified topics
        4. If a location has no relevant law changes, return empty lawChanges array for that location
        5. DO NOT include any jurisdictions not explicitly listed above
        6. DO NOT include any industries not listed above
        7. DO NOT include any topics not listed above

        Output Format: Group by the EXACT locations specified above. Use proper location names (not slugs).

        If no law changes match ALL the criteria (location + industry + topic), return empty arrays."""
        )
        messages = [
            system_message,
            HumanMessage(content=f"text: {vectors_as_string}")
        ]
        structured_llm = llm.with_structured_output(LawChangeListByLocation)
        response = structured_llm.invoke(messages)
        return response
    
    
    def extract_news(self, vectors: list):
        vector_list = []

        for match in vectors.get("matches", []):
            text = match["metadata"].get("chunk_text") or match["metadata"].get("text", "")
            metadata = match["metadata"]
            vector_list.append({
                "text": text,
                "metadata": metadata
            })
        vectors_as_string = "\n\n---\n\n".join([
            f"Text: {v['text']}\n\nMetadata: {v['metadata']}" for v in vector_list
        ])
        llm=AzureChatOpenAI(
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
            azure_deployment=os.getenv('gpt-4o-mini'),
            api_key=os.getenv('AZURE_OPENAI_KEY'),
            api_version="2024-12-01-preview",
            model_name="gpt-4o-mini",
            openai_api_type="azure",
            streaming=True,
            callbacks=[StreamingStdOutCallbackHandler()]
        )
        system_message = SystemMessage(
            content=f"""You are a legal news analyst. Analyze the provided documents and extract at least 4 unique structured news items about
            changes in employment law or HR regulations.""")

        messages = [
            system_message,
            HumanMessage(content=f"text: {vectors_as_string}")
        ]
        structured_llm = llm.with_structured_output(NewsList)
        response = structured_llm.invoke(messages)
        return response
    
    def extract_court_decisions(self, vectors: list, location_slugs=None, industry_slugs=None, topic_slugs=None):
        vector_list = []
        for match in vectors.get("matches", []):
            text = match["metadata"].get("chunk_text") or match["metadata"].get("text", "")
            metadata = match["metadata"]
            vector_list.append({
                "text": text,
                "metadata": metadata
            })
        vectors_as_string = "\n\n---\n\n".join([
            f"Text: {v['text']}\n\nMetadata: {v['metadata']}" for v in vector_list
        ])

        llm = AzureChatOpenAI(
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
            azure_deployment=os.getenv('gpt-4o-mini'),
            api_key=os.getenv('AZURE_OPENAI_KEY'),
            api_version="2024-12-01-preview",
            model_name="gpt-4o-mini",
            openai_api_type="azure",
            streaming=True,
            callbacks=[StreamingStdOutCallbackHandler()]
        )
        location_slugs = location_slugs or []
        industry_slugs = industry_slugs or []
        topic_slugs = topic_slugs or []

        location_names = [slug.replace('-', ' ').title() for slug in location_slugs]
        industry_names = [slug.replace('-', ' ').title() for slug in industry_slugs]
        topic_names = [slug.replace('-', ' ').title() for slug in topic_slugs]

        system_message = SystemMessage(
            content=(
                "You are a legal research assistant. Analyze the provided text and extract ONLY actual court decisions.\n\n"
                "**What qualifies as a court decision:**\n"
                "- Includes a specific case name (e.g., 'Groff v. DeJoy')\n"
                "- Includes a ruling or judgment that impacts employers\n\n"
                "**What to exclude:**\n"
                "- Generic laws, EEOC guidance, or policy summaries\n"
                "- Articles or blog posts without an actual case name\n\n"
                "**Filtering:**\n"
                f"Required locations: {', '.join(location_names) if location_names else 'None'}\n"
                f"Required industries: {', '.join(industry_names) if industry_names else 'None'}\n"
                f"Required topics: {', '.join(topic_names) if topic_names else 'None'}\n\n"
                "**Output format:**\n"
                "Return a valid JSON object with this structure:\n"
                "{\n"
                "  \"courtDecisions\": [\n" 
                "    {\n"
                "      \"title\": \"Case name (e.g., Groff v. DeJoy)\",\n"
                "      \"description\": \"Summary of the decision and its impact on employers\",\n"
                "      \"sourceUrl\": \"The URL from the text or an empty string if not available\"\n"
                "    }\n"
                "  ]\n"
                "}\n"
                "If there are no valid cases, return: {\"courtDecisions\": []}"
            )
        )


        messages = [
            system_message,
            HumanMessage(content=f"text: {vectors_as_string}")
        ]
        structured_llm = llm.with_structured_output(CourtDecisionList)
        response = structured_llm.invoke(messages)
        return response
    def extract_legal_calendar(self, vectors: list):
            vector_list = []

            for match in vectors.get("matches", []):
                text = match["metadata"].get("chunk_text") or match["metadata"].get("text", "")
                metadata = match["metadata"]
                vector_list.append({
                    "text": text,
                    "metadata": metadata
                })

            vectors_as_string = "\n\n---\n\n".join([
                f"Text: {v['text']}\n\nMetadata: {v['metadata']}" for v in vector_list
            ])

            llm = AzureChatOpenAI(
                azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
                azure_deployment=os.getenv('gpt-4o-mini'),
                api_key=os.getenv('AZURE_OPENAI_KEY'),
                api_version="2024-12-01-preview",
                model_name="gpt-4o-mini",
                openai_api_type="azure",
                streaming=True,
                callbacks=[StreamingStdOutCallbackHandler()]
            )

            system_message = SystemMessage(
                content="""You are a legal news analyst. Analyze the following text and extract unique legal calendar events related to changes in employment laws or HR regulation
                Return a structured list of legal calendar events."""
            )

            messages = [
                system_message,
                HumanMessage(content=f"text: {vectors_as_string}")
            ]

            structured_llm = llm.with_structured_output(LegalCalendar)
            response = structured_llm.invoke(messages)
            return response
